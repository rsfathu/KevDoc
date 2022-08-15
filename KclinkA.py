import pandas
from sqlalchemy import create_engine
import sqlite3
import random as r
import numpy
import pprint
from sqlite3.dbapi2 import SQLITE_SELECT
from docx import Document
import os.path

x=['Timmy','did not','shopping','yesterday'] # this simulates the results that would be gotten from F2Rest
reporttype = 'praise_letter' # this is the report type that would be taken from user input
name_to_save = 'Xyz' # this is the saved name that would be taken from user input

database = sqlite3.connect('DocXPanda.db')
cursor = database.cursor()


# XXX selects and return template to be used, exactly like F2Rest's initiateprogram() and retrieve() combined XXX
def selecttemplate(templatechosen):

    template_no = templatechosen

    # Using Pandas to get the latest data from google sheets table with templates
    sheet_id2 = "1IB3HAKdWgcybrBWta61WYFTG3MLh40iWytgz_SVAl-4"
    # df2 stands for datafile2
    df2 = pandas.read_csv(f"https://docs.google.com/spreadsheets/d/{sheet_id2}/export?format=csv") # retrieve relevant columns specifically
    idcolumn2 = df2[[template_no
                     ]]  #extract row where template is stored

    engine = create_engine('sqlite:///DocXPanda.db', echo=True)
    sqlite_connection = engine.connect()
    # Creating the table with headers
    idcolumn2.to_sql("DocX_information", sqlite_connection, if_exists="replace")

    executedcode = " SELECT {} FROM 'DocX_information' ".format(str(templatechosen))
    cursor.execute(executedcode)
    getdata = cursor.fetchall() # this extracts the template to be used

    return getdata

# XXX Subs in variables into the template to form the report XXX
def sub_in_variables(variables,template):

    #for i in range(0,len(variables)+1):
    #    print(i)
    #    final_print = [str(template[0]),str(template[1]).replace('=|{}|='.format(str(i-1)), variables[i-1])]

    midprint = str(template[1]).replace('=|1|=', variables[0])
    midprint1 = midprint.replace('=|2|=', variables[1])
    midprint2 = midprint1.replace('=|3|=', variables[2])
    midprint3 = midprint2.replace('=|4|=', variables[3])

    midprint4 = midprint3.replace("('","")
    midprint5 = midprint4.replace("',)", "")

    final_print = [template[0],midprint5, template[2]]

    return final_print

# XXX this creates the document and saves it XXX
def print_doc(savename,content):

    document = Document()
    content2 = content[1].split('<p>')

    k = content[2] # converts title type into an integer
    k1 = int(k[0])

    document.add_heading(content[0], k1) # adds heading and paragraphes
    for i in range(0,len(content2)):
        document.add_paragraph(content2[i])
    document.save('C:/Users/KC Chan/Desktop/'+ savename + str(1) + '.docx')

print_doc(name_to_save,sub_in_variables(x,selecttemplate(reporttype)))
