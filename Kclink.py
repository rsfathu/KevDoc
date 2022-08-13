from docx import Document
from docx.shared import Inches
from F2Rest import *

document = Document()

#this is to add a title
h1 = document.add_heading('Title ' + str(1), 0)

#a, b, c are random numbers from 0 to 4        
a = r.randint(0,4)
b = r.randint(0,4)
c = r.randint(0,4)
#d is a sentence that is constructed from three random words and "did not"
d = str((str(subjectlst[a]) + str(modallst[0]) + str(verblst[b]) + str(adverbiallst[c])))

#this is to add d into a normal paragraph
p1 = document.add_paragraph(d)

#this is to save the word document
document.save('KevDoc ' + str(1) + '.docx')

#.clear is to clear ONLY the syntax; not the style of the heading/paragraph
h1.clear()
p1.clear()

i = 1
while i != 4:
    i += 1
    h1.add_run('Title ' + str(i))
        
    a = r.randint(0,4)
    b = r.randint(0,4)
    c = r.randint(0,4)
    d = str((str(subjectlst[a]) + str(modallst[0]) + str(verblst[b]) + str(adverbiallst[c])))

    p1.add_run(d)

    document.save('KevDoc ' + str(i) + '.docx')

    h1.clear()
    p1.clear()


"""
#This is the actual version I got from KC

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

#document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')

"""